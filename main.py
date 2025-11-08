import customtkinter as ctk
from tkinter import filedialog, scrolledtext, messagebox
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from PyPDF2 import PdfReader
from openai import OpenAI
import cv2, mediapipe as mp, threading, math, queue, json, os
import sounddevice as sd
import vosk
from PIL import Image, ImageTk

# ---------------- Hugging Face ----------------
HF_TOKEN = "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"
client = OpenAI(base_url="https://router.huggingface.co/v1", api_key=HF_TOKEN)

# ---------------- Global ----------------
camera_active = False
points = []
marker_color = "yellow"
file_content_global = ""
highlight_tags = set()
scroll_speed = 3
camera_frame = None

# ---------------- MediaPipe ----------------
mp_hands = mp.solutions.hands
mp_drawing = mp.solutions.drawing_utils

# ---------------- File Reading ----------------
def read_file(path):
    text = ""
    if path.endswith(".txt"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    elif path.endswith(".docx"):
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif path.endswith(".pdf"):
        reader = PdfReader(path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    else:
        messagebox.showerror("Unsupported File", "Use .txt, .docx, or .pdf")
    return text

# ---------------- AI Interaction ----------------
def ask_ai():
    user_input = user_text.get("1.0", "end").strip()
    if not user_input:
        return
    chat_box.configure(state="normal")
    chat_box.insert("end", f"🧑 You: {user_input}\n")
    chat_box.configure(state="disabled")
    chat_box.see("end")
    try:
        prompt = f"File content:\n{file_content_global}\n\nQuestion: {user_input}\nAnswer:"
        completion = client.chat.completions.create(
            model="moonshotai/Kimi-K2-Instruct-0905",
            messages=[{"role": "user", "content": prompt}]
        )
        ai_response = completion.choices[0].message.content
    except Exception as e:
        ai_response = f"Error: {e}"
    chat_box.configure(state="normal")
    chat_box.insert("end", f"🤖 AI: {ai_response}\n\n")
    chat_box.configure(state="disabled")
    chat_box.see("end")
    user_text.delete("1.0", "end")

# ---------------- Gesture Utils ----------------
def count_fingers(hand_landmarks):
    tips = [mp_hands.HandLandmark.THUMB_TIP,
            mp_hands.HandLandmark.INDEX_FINGER_TIP,
            mp_hands.HandLandmark.MIDDLE_FINGER_TIP,
            mp_hands.HandLandmark.RING_FINGER_TIP,
            mp_hands.HandLandmark.PINKY_TIP]
    cnt = 0
    for i, tip in enumerate(tips):
        if i == 0:
            if hand_landmarks.landmark[tip].x < hand_landmarks.landmark[tip - 2].x:
                cnt += 1
        else:
            if hand_landmarks.landmark[tip].y < hand_landmarks.landmark[tip - 2].y:
                cnt += 1
    return cnt

def is_circle(points):
    if len(points) < 30: return False
    x_vals = [p[0] for p in points]
    y_vals = [p[1] for p in points]
    width, height = max(x_vals) - min(x_vals), max(y_vals) - min(y_vals)
    aspect_ratio = width / height if height != 0 else 0
    dx = points[0][0] - points[-1][0]
    dy = points[0][1] - points[-1][1]
    distance = math.hypot(dx, dy)
    return 0.7 <= aspect_ratio <= 1.3 and distance < 50

# ---------------- Highlight ----------------
def highlight_word(x, y, color):
    try:
        index = file_text_box.index(f"@{x},{y}")
    except:
        return
    line_num = index.split('.')[0]
    line_text = file_text_box.get(f"{line_num}.0", f"{line_num}.end")
    char_idx = min(int(index.split('.')[1]), len(line_text) - 1)
    left, right = char_idx, char_idx
    while left > 0 and line_text[left - 1] not in [" ", "\n"]: left -= 1
    while right < len(line_text) - 1 and line_text[right + 1] not in [" ", "\n"]: right += 1
    start = f"{line_num}.{left}"
    end = f"{line_num}.{right + 1}"
    tag = f"hl_{len(highlight_tags)}"
    highlight_tags.add(tag)
    file_text_box.tag_add(tag, start, end)
    file_text_box.tag_config(tag, background=color)
    file_text_box.see(start)

# ---------------- Scroll ----------------
def scroll_text(fingers):
    if fingers == 5:
        file_text_box.yview_scroll(scroll_speed, "units")
    elif fingers == 2:
        file_text_box.yview_scroll(-scroll_speed, "units")

# ---------------- Camera ----------------
def detect_gestures():
    global camera_active, points, camera_frame
    cap = cv2.VideoCapture(0)
    hands = mp_hands.Hands(min_detection_confidence=0.7, min_tracking_confidence=0.7)
    points = []
    while camera_active:
        ret, frame = cap.read()
        if not ret: break
        frame = cv2.flip(frame, 1)
        rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        results = hands.process(rgb)
        if results.multi_hand_landmarks:
            hand = results.multi_hand_landmarks[0]
            mp_drawing.draw_landmarks(frame, hand, mp_hands.HAND_CONNECTIONS)
            fingers = count_fingers(hand)
            x = int(hand.landmark[mp_hands.HandLandmark.INDEX_FINGER_TIP].x * frame.shape[1])
            y = int(hand.landmark[mp_hands.HandLandmark.INDEX_FINGER_TIP].y * frame.shape[0])
            if fingers == 1: highlight_word(x, y, marker_color)
            scroll_text(fingers)
            points.append((x, y))
            if len(points) > 50: points.pop(0)
            if is_circle(points):
                x_vals = [p[0] for p in points]
                y_vals = [p[1] for p in points]
                start = file_text_box.index(f"@{min(x_vals)},{min(y_vals)}")
                end = file_text_box.index(f"@{max(x_vals)},{max(y_vals)}")
                copied = file_text_box.get(start, end)
                user_text.delete("1.0", "end")
                user_text.insert("1.0", copied)
                points = []
        overlay = cv2.resize(frame, (300, 200))
        overlay = cv2.cvtColor(overlay, cv2.COLOR_BGR2RGB)
        overlay = Image.fromarray(overlay)
        camera_frame = ImageTk.PhotoImage(overlay)
        camera_label.configure(image=camera_frame)
    cap.release()

def start_camera_thread():
    global camera_active
    if not camera_active:
        camera_active = True
        threading.Thread(target=detect_gestures, daemon=True).start()

# ---------------- Voice ----------------
def listen_voice():
    global marker_color

    # Change this if your model is elsewhere
    model_path =r"C:\Users\anjum\OneDrive\Desktop\AI FILE VIEWER\vosk-model-small-en-us-0.15"


    # 1️⃣ Check if folder exists
    if not os.path.isdir(model_path):
        print(f"❌ Vosk model folder not found at:\n{model_path}")
        print("➡ Please extract the model ZIP file and ensure it contains folders like 'am', 'conf', 'graph'.")
        return

    # 2️⃣ Try loading the model
    try:
        model = vosk.Model(model_path)
        print("✅ Vosk model loaded successfully!")
    except Exception as e:
        print("❌ Failed to load Vosk model:", e)
        print("➡ Try re-downloading from: https://alphacephei.com/vosk/models")
        return

    # 3️⃣ Start voice recognition
    q = queue.Queue()

    def callback(indata, frames, time, status):
        q.put(bytes(indata))

    with sd.RawInputStream(samplerate=16000, blocksize=8000, dtype='int16',
                           channels=1, callback=callback):
        rec = vosk.KaldiRecognizer(model, 16000)
        print("🎙 Voice recognition started! Say 'camera on', 'camera off', or a color (red, blue, green, etc).")
        while True:
            data = q.get()
            if rec.AcceptWaveform(data):
                result = json.loads(rec.Result())
                if 'text' in result:
                    cmd = result['text'].lower()
                    if cmd:
                        print("🗣 Heard:", cmd)
                    if "camera on" in cmd:
                        start_camera_thread()
                    elif "camera off" in cmd:
                        global camera_active
                        camera_active = False
                    elif cmd in ["red", "green", "blue", "yellow", "cyan", "magenta"]:
                        marker_color = cmd
                        print("✏ Marker color changed to", marker_color)

# ---------------- Download DOCX ----------------
color_map = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "red": WD_COLOR_INDEX.RED,
    "blue": WD_COLOR_INDEX.BLUE,
    "cyan": WD_COLOR_INDEX.TURQUOISE,
    "magenta": WD_COLOR_INDEX.PINK
}

def download_content():
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if path:
        doc = Document()
        doc.add_heading("Full Text with Highlights", 0)
        lines = file_text_box.get("1.0", "end-1c").split("\n")
        for i, line in enumerate(lines):
            p = doc.add_paragraph()
            idx = f"{i + 1}.0"
            for char_idx, char in enumerate(line):
                included = False
                for tag in highlight_tags:
                    ranges = file_text_box.tag_ranges(tag)
                    for j in range(0, len(ranges), 2):
                        start, end = ranges[j], ranges[j + 1]
                        if file_text_box.compare(f"{idx}+{char_idx}c", ">=", start) and \
                           file_text_box.compare(f"{idx}+{char_idx}c", "<", end):
                            run = p.add_run(char)
                            color_name = file_text_box.tag_cget(tag, "background")
                            run.font.highlight_color = color_map.get(color_name, WD_COLOR_INDEX.YELLOW)
                            included = True
                if not included:
                    p.add_run(char)
        doc.add_page_break()
        doc.add_heading("Chat History", level=1)
        doc.add_paragraph(chat_box.get("1.0", "end-1c"))
        doc.save(path)
        messagebox.showinfo("Saved", f"File saved at:\n{path}")

# ---------------- Display ----------------
def display_file(text):
    global file_text_box, user_text, chat_box, camera_label, file_content_global
    file_content_global = text
    win = ctk.CTkToplevel(root)
    win.geometry("1500x900")
    win.title("Professional AI File Viewer")

    left = ctk.CTkFrame(win, width=1000)
    left.pack(side="left", fill="both", expand=True, padx=10, pady=10)
    file_text_box = scrolledtext.ScrolledText(left, font=("Calibri", 14))
    file_text_box.pack(expand=True, fill="both")
    file_text_box.insert("1.0", text)
    file_text_box.config(state="normal")

    right = ctk.CTkFrame(win, width=450)
    right.pack(side="right", fill="y", padx=10, pady=10)
    ctk.CTkLabel(right, text="🤖 AI Explainer", font=("Calibri", 16, "bold")).pack(pady=10)
    user_text = ctk.CTkTextbox(right, height=80, font=("Calibri", 12))
    user_text.pack(fill="x", padx=10, pady=5)
    ctk.CTkButton(right, text="Send", command=ask_ai).pack(pady=5)
    ctk.CTkLabel(right, text="Chat History", font=("Calibri", 12, "bold")).pack(pady=(10, 0))
    chat_box = ctk.CTkTextbox(right, height=400, font=("Calibri", 12))
    chat_box.pack(fill="both", expand=True, padx=10, pady=5)
    chat_box.configure(state="disabled")
    ctk.CTkButton(right, text="Download DOCX", width=150,
                  fg_color="#28a745", command=download_content).pack(pady=10)
    camera_label = ctk.CTkLabel(right, text="Camera Loading...")
    camera_label.pack(pady=20)

# ---------------- Open File ----------------
def open_file():
    path = filedialog.askopenfilename(title="Select file",
                                      filetypes=[("Text", ".txt"), ("Word", ".docx"),
                                                 ("PDF", ".pdf"), ("All", ".")])
    if path:
        content = read_file(path)
        if content.strip():
            display_file(content)
        else:
            messagebox.showinfo("Empty", "File has no readable text.")

# ---------------- Main Window ----------------
ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("blue")

root = ctk.CTk()
root.geometry("550x250")
root.title("Modern AI File Viewer")

ctk.CTkLabel(root, text="Choose a file to open", font=("Calibri", 14)).pack(pady=30)
ctk.CTkButton(root, text="Open File", width=150, command=open_file).pack(pady=20)

threading.Thread(target=listen_voice, daemon=True).start()

root.mainloop()